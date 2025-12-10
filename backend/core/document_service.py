from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from typing import Dict, Any
from datetime import datetime
from fastapi.responses import FileResponse
from fastapi import HTTPException, status
import tempfile
from typing import Dict, Any
from io import BytesIO
import logging

logger = logging.getLogger(__name__)

def generate_document(proposal_data: Dict[str, Any], format: str = "docx") -> bytes:
    """
    Genera un documento con la propuesta comercial.
    
    Args:
        proposal_data: Datos del análisis de la propuesta
        format: Formato del documento ("docx" o "pdf")
        
    Returns:
        Documento generado como bytes
        
    Raises:
        HTTPException 400: Si el formato no es soportado
        HTTPException 500: Si hay error en la generación
    """
    try:
        if format.lower() == "docx":
            return _generate_docx(proposal_data)
        elif format.lower() == "pdf":
            return _generate_pdf(proposal_data)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Formato no soportado: {format}. Use 'docx' o 'pdf'."
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error al generar documento ({format}): {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error al generar el documento: {str(e)}"
        )

def _generate_docx(proposal_data: Dict[str, Any]) -> bytes:
    """
    Genera un docuemto de Word basado en los datos de la propuesta proporcionados.

    Args:
        proposal_data (Dict[str, Any]): La información puede incluir título, descripción, objetivos, cronograma, presupuesto, etc.

    Returns:
        str: La ruta del archivo al documento de Word generado.
    """
    # Crear documento Word
    doc = Document()

    # ============ CONFIGURAR TIPOGRAFÍA GLOBAL ============
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(50, 50, 50)  # Gris oscuro elegante

    # ============ ESTILOS PERSONALIZADOS ============

    # Título Corporativo
    title_style = doc.styles.add_style('CorporateTitle', 1)
    title_style.font.name = 'Calibri Light'
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri Light')
    title_style.font.size = Pt(32)
    title_style.font.color.rgb = RGBColor(30, 30, 30)

    # Encabezados de sección
    header_style = doc.styles.add_style('CorporateHeader', 1)
    header_style.font.name = 'Segoe UI Semibold'
    header_style.font.size = Pt(18)
    header_style.font.color.rgb = RGBColor(40, 40, 120)  # Azul corporativo

    # Encabezados de subsección
    subheader_style = doc.styles.add_style('CorporateSubHeader', 1)
    subheader_style.font.name = 'Segoe UI'
    subheader_style.font.size = Pt(14)
    subheader_style.font.color.rgb = RGBColor(70, 70, 70)

    # ============ PORTADA ============
    title = doc.add_paragraph("PROPUESTA TÉCNICA", style="CorporateTitle")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("\n")
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Cliente: {proposal_data.get('cliente', 'No especificado')}")
    doc.add_paragraph("\nDocumento confidencial — Uso exclusivo del cliente.")

    doc.add_page_break()

    # ============ SECCIÓN: INTRODUCCIÓN ============
    doc.add_paragraph("1. Introducción", style="CorporateHeader")
    doc.add_paragraph(
        "El presente documento presenta la propuesta técnica elaborada por nuestro equipo, "
        "definida con base en los lineamientos y requerimientos compartidos por el cliente."
    )

    # ============ SECCIÓN: OBJETIVO ============
    doc.add_paragraph("2. Objetivo del Proyecto", style="CorporateHeader")
    doc.add_paragraph(
        proposal_data.get(
            "objetivo_proyecto",
            "Presentar una solución técnica integral que responda a las necesidades identificadas."
        )
    )

    # ============ SECCIÓN: INFORMACIÓN GENERAL ============
    doc.add_paragraph("3. Información General del Proyecto", style="CorporateHeader")
    doc.add_paragraph(f"Fecha estimada de entrega: {proposal_data.get('fecha_entrega', 'No especificada')}")
    doc.add_paragraph(f"Cliente: {proposal_data.get('cliente', 'No especificado')}")

    # ============ SECCIÓN: ALCANCE ECONÓMICO ============
    doc.add_paragraph("4. Alcance Económico", style="CorporateHeader")
    alcance = proposal_data.get('alcance_economico', {})
    doc.add_paragraph(f"Presupuesto Estimado: {alcance.get('presupuesto', 'No especificado')}")
    doc.add_paragraph(f"Moneda: {alcance.get('moneda', 'No especificada')}")

    doc.add_paragraph(
        "Los valores expresados están sujetos a validación final y pueden ajustarse de acuerdo "
        "con requerimientos adicionales o cambios en el alcance del proyecto."
    )

    # ============ SECCIÓN: TECNOLOGÍAS ============
    doc.add_paragraph("5. Tecnologías Propuestas", style="CorporateHeader")

    tecnologias = proposal_data.get('tecnologias_requeridas', [])
    if tecnologias:
        p = doc.add_paragraph("Las tecnologías sugeridas son las siguientes:")
        for tech in tecnologias:
            doc.add_paragraph(f"• {tech}", style="List Bullet")
    else:
        doc.add_paragraph("No se especificaron tecnologías requeridas.")

    # ============ SECCIÓN: PREGUNTAS ============
    doc.add_paragraph("6. Preguntas para Aclaración", style="CorporateHeader")
    preguntas = proposal_data.get('preguntas_sugeridas', [])

    if preguntas:
        for pregunta in preguntas:
            doc.add_paragraph(f"• {pregunta}", style="List Bullet")
    else:
        doc.add_paragraph("No se registraron preguntas.")

    # ============ SECCIÓN: EQUIPO ============
    doc.add_paragraph("7. Equipo Propuesto", style="CorporateHeader")
    equipo = proposal_data.get('equipo_sugerido', [])

    for miembro in equipo:
        doc.add_paragraph(miembro.get('nombre', 'Sin nombre'), style="CorporateSubHeader")
        doc.add_paragraph(f"Rol: {miembro.get('rol', 'No especificado')}")
        doc.add_paragraph(f"Experiencia: {miembro.get('experiencia', 'No especificada')}")

        skills = miembro.get('skills', [])
        if skills:
            doc.add_paragraph("Habilidades Clave:")
            for skill in skills:
                doc.add_paragraph(f"• {skill}", style="List Bullet")
    
    # Guardar documento temporalmente
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        tmp_path = tmp_file.name
        doc.save(tmp_path)
        
    file_name = f"Propuesta_{proposal_data.get('cliente', 'Cliente')}.docx"
    
    # Guardar en memoria
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Devolver documento
    return output.getvalue()

def _generate_pdf(proposal_data: Dict[str, Any]) -> bytes:
    """Genera documento PDF usando ReportLab"""
    
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Estilos personalizados
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=30,
        alignment=1  # Center
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#1e3a8a'),
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Título
    story.append(Paragraph("Propuesta Comercial", title_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Cliente
    story.append(Paragraph("Cliente", heading_style))
    story.append(Paragraph(proposal_data.get("cliente", "N/A"), styles['Normal']))
    story.append(Spacer(1, 0.15*inch))
    
    # Alcance Económico
    story.append(Paragraph("Alcance Económico", heading_style))
    alcance = proposal_data.get("alcance_economico", {})
    presupuesto_text = f"<b>Presupuesto:</b> {alcance.get('presupuesto', 'N/A')} {alcance.get('moneda', 'USD')}"
    story.append(Paragraph(presupuesto_text, styles['Normal']))
    story.append(Spacer(1, 0.15*inch))
    
    # Fecha Entrega
    story.append(Paragraph("Fecha de Entrega", heading_style))
    story.append(Paragraph(proposal_data.get("fecha_entrega", "N/A"), styles['Normal']))
    story.append(Spacer(1, 0.15*inch))
    
    #Objetivo General
    story.append(Paragraph("Objetivo General", heading_style))
    story.append(Paragraph(proposal_data.get("objetivo_general", "N/A"), styles['Normal']))
    story.append(Spacer(1, 0.15*inch))
    
    # Preguntas Sugeridas
    preguntas = proposal_data.get("preguntas_sugeridas", [])
    if preguntas:
        story.append(Paragraph("Preguntas Sugeridas", heading_style))
        for pregunta in preguntas:
            story.append(Paragraph(f"• {pregunta}", styles['Normal']))
        story.append(Spacer(1, 0.15*inch))
    
    # Equipo Sugerido
    equipo = proposal_data.get("equipo_sugerido", [])
    if equipo:
        story.append(Paragraph("Equipo Sugerido", heading_style))
        for miembro in equipo:
            nombre = miembro.get('nombre', 'N/A')
            rol = miembro.get('rol', 'N/A')
            exp = miembro.get('experiencia', 'N/A')
            skills = ', '.join(miembro.get('skills', []))
            
            miembro_text = f"<b>{nombre}</b><br/>Rol: {rol}<br/>Experiencia: {exp}<br/>Skills: {skills}"
            story.append(Paragraph(miembro_text, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        story.append(Spacer(1, 0.15*inch))
    
    # Tecnologías
    techs = proposal_data.get("tecnologias_requeridas", [])
    if techs:
        story.append(Paragraph("Tecnologías Requeridas", heading_style))
        story.append(Paragraph(", ".join(techs), styles['Normal']))
    
    # Construir PDF
    doc.build(story)
    output.seek(0)
    
    # Devolver documento
    return output.getvalue()

