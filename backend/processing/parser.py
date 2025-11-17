import os
from pathlib import Path
import pandas as pd
import PyPDF2
import docx
import openpyxl # Necesario para que pandas lea .xlsx
import mimetypes

def extract_text_from_file(file_path: Path) -> str:
    """
    Extrae texto de un archivo (PDF, DOCX, XLSX, CSV, TXT).
    Usa detección por extensión como fallback si mimetypes falla.
    """
    # Intentar detectar tipo MIME
    file_type = mimetypes.guess_type(file_path)[0]
    
    # Fallback: detectar por extensión si mimetypes falla
    file_extension = Path(file_path).suffix.lower()
    
    if not file_type:
        extension_to_mimetype = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.csv': 'text/csv',
            '.txt': 'text/plain',
        }
        file_type = extension_to_mimetype.get(file_extension)
        print(f"PARSER: MIME type no detectado, usando extensión: {file_extension} -> {file_type}")
    
    print(f"PARSER: Extrayendo texto de {file_path} (Tipo: {file_type}, Extensión: {file_extension})")
    
    full_text = ""
    
    try:
        if file_type == "application/pdf" or file_extension == '.pdf':
            # Intentar con PyPDF2 primero
            try:
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            full_text += page_text + "\n"
                
                # Si PyPDF2 solo extrajo espacios/caracteres vacíos
                if len(full_text.strip()) < 10:
                    print(f"PARSER: PyPDF2 extrajo muy poco texto ({len(full_text.strip())} chars), intentando con pdfplumber...")
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                full_text += page_text + "\n"
            except Exception as pdf_error:
                print(f"PARSER: Error con PyPDF2, intentando pdfplumber: {pdf_error}")
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            full_text += page_text + "\n"
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_extension == '.docx':
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                full_text += para.text + "\n"
            # También extraer de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + " "
                    full_text += "\n"
        
        elif file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" or file_extension == '.xlsx':
            df = pd.read_excel(file_path)
            full_text = df.to_string() # Convierte el dataframe entero a texto
        
        elif file_type == "text/csv" or file_extension == '.csv':
            df = pd.read_csv(file_path)
            full_text = df.to_string()
            
        elif file_type == "text/plain" or file_extension == '.txt':
            with open(file_path, "r", encoding="utf-8") as f:
                full_text = f.read()
        
        else:
            print(f"PARSER: Tipo de archivo '{file_type}' no soportado.")
            raise ValueError(f"Tipo de archivo no soportado: {file_type}")
        
        # Validar que se extrajo texto útil
        if len(full_text.strip()) < 10:
            print(f"PARSER: ADVERTENCIA - Texto extraído muy corto ({len(full_text.strip())} caracteres). El archivo puede estar vacío o ser una imagen escaneada.")
        
        print(f"PARSER: Texto extraído exitosamente (Longitud: {len(full_text.strip())} caracteres).")
        return full_text

    except Exception as e:
        print(f"PARSER: Error al extraer texto de {file_path}: {e}")
        raise