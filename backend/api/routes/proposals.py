"""
Endpoints para análisis y generación de propuestas comerciales.

Este módulo proporciona endpoints para:
- Analizar documentos RFP con IA
- Generar propuestas en formato Word
"""

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
from typing import Dict, Any
from api.service.impl.proposals_service_impl import ProposalsServiceImpl
from api.service.proposals_service import ProposalsService
from models import database
from models.user import User
from core.auth import get_current_active_user
from core.llm_service import get_provider
import logging
import tempfile

# Configurar logger
logger = logging.getLogger(__name__)

router = APIRouter()

def get_proposal_service() -> ProposalsService:
    """Función de dependencia para obtener la instancia del servicio."""
    return ProposalsServiceImpl()



@router.post("/proposals/analyze", summary="Actualizado")
async def analyze_proposal(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_active_user),
    service: ProposalsService = Depends(get_proposal_service) 
):
    """Delega toda la lógica de validación, extracción y análisis al servicio."""
    return await service.analyze(file=file)


    

@router.post(
    "/proposals/generate",
    summary="Generar documento de propuesta",
    description="Genera un documento Word con la propuesta comercial"
)
async def generate_proposal_document(
    proposal_data: Dict[str, Any],
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(database.get_db)
):
    """
    Genera un documento Word con la propuesta comercial.
    
    Args:
        proposal_data: Datos del análisis de la propuesta
        current_user: Usuario autenticado
        db: Sesión de base de datos
        
    Returns:
        Documento Word generado
        
    Raises:
        HTTPException 500: Si hay error en la generación
    """
    
    try:
        # Crear documento Word
        doc = Document()

        # ============ CONFIGURAR TIPOGRAFÍA GLOBAL ============
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(50, 50, 50)  # Gris oscuro elegante

        # ============ ESTILOS PERSONALIZADOS ============

        # Título Corporativo
        title_style = doc.styles.add_style('CorporateTitle', 1)
        title_style.font.name = 'Calibri Light'
        title_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri Light')
        title_style.font.size = Pt(32)
        title_style.font.color.rgb = RGBColor(30, 30, 30)

        # Encabezados de sección
        header_style = doc.styles.add_style('CorporateHeader', 1)
        header_style.font.name = 'Segoe UI Semibold'
        header_style.font.size = Pt(18)
        header_style.font.color.rgb = RGBColor(40, 40, 120)  # Azul corporativo

        # Encabezados de subsección
        subheader_style = doc.styles.add_style('CorporateSubHeader', 1)
        subheader_style.font.name = 'Segoe UI'
        subheader_style.font.size = Pt(14)
        subheader_style.font.color.rgb = RGBColor(70, 70, 70)

        # ============ PORTADA ============
        title = doc.add_paragraph("PROPUESTA TÉCNICA", style="CorporateTitle")
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("\n")
        doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d')}")
        doc.add_paragraph(f"Cliente: {proposal_data.get('cliente', 'No especificado')}")
        doc.add_paragraph("\nDocumento confidencial — Uso exclusivo del cliente.")

        doc.add_page_break()

        # ============ SECCIÓN: INTRODUCCIÓN ============
        doc.add_paragraph("1. Introducción", style="CorporateHeader")
        doc.add_paragraph(
            "El presente documento presenta la propuesta técnica elaborada por nuestro equipo, "
            "definida con base en los lineamientos y requerimientos compartidos por el cliente."
        )

        # ============ SECCIÓN: OBJETIVO ============
        doc.add_paragraph("2. Objetivo del Proyecto", style="CorporateHeader")
        doc.add_paragraph(
            proposal_data.get(
                "objetivo_proyecto",
                "Presentar una solución técnica integral que responda a las necesidades identificadas."
            )
        )

        # ============ SECCIÓN: INFORMACIÓN GENERAL ============
        doc.add_paragraph("3. Información General del Proyecto", style="CorporateHeader")
        doc.add_paragraph(f"Fecha estimada de entrega: {proposal_data.get('fecha_entrega', 'No especificada')}")
        doc.add_paragraph(f"Cliente: {proposal_data.get('cliente', 'No especificado')}")

        # ============ SECCIÓN: ALCANCE ECONÓMICO ============
        doc.add_paragraph("4. Alcance Económico", style="CorporateHeader")
        alcance = proposal_data.get('alcance_economico', {})
        doc.add_paragraph(f"Presupuesto Estimado: {alcance.get('presupuesto', 'No especificado')}")
        doc.add_paragraph(f"Moneda: {alcance.get('moneda', 'No especificada')}")

        doc.add_paragraph(
            "Los valores expresados están sujetos a validación final y pueden ajustarse de acuerdo "
            "con requerimientos adicionales o cambios en el alcance del proyecto."
        )

        # ============ SECCIÓN: TECNOLOGÍAS ============
        doc.add_paragraph("5. Tecnologías Propuestas", style="CorporateHeader")

        tecnologias = proposal_data.get('tecnologias_requeridas', [])
        if tecnologias:
            p = doc.add_paragraph("Las tecnologías sugeridas son las siguientes:")
            for tech in tecnologias:
                doc.add_paragraph(f"• {tech}", style="List Bullet")
        else:
            doc.add_paragraph("No se especificaron tecnologías requeridas.")

        # ============ SECCIÓN: RIESGOS ============
        doc.add_paragraph("6. Riesgos Identificados", style="CorporateHeader")
        riesgos = proposal_data.get('riesgos_detectados', [])

        if riesgos:
            for riesgo in riesgos:
                doc.add_paragraph(f"• {riesgo}", style="List Bullet")
        else:
            doc.add_paragraph("No se identificaron riesgos relevantes.")

        # ============ SECCIÓN: PREGUNTAS ============
        doc.add_paragraph("7. Preguntas para Aclaración", style="CorporateHeader")
        preguntas = proposal_data.get('preguntas_sugeridas', [])

        if preguntas:
            for pregunta in preguntas:
                doc.add_paragraph(f"• {pregunta}", style="List Bullet")
        else:
            doc.add_paragraph("No se registraron preguntas.")

        # ============ SECCIÓN: EQUIPO ============
        doc.add_paragraph("8. Equipo Propuesto", style="CorporateHeader")
        equipo = proposal_data.get('equipo_sugerido', [])

        for miembro in equipo:
            doc.add_paragraph(miembro.get('nombre', 'Sin nombre'), style="CorporateSubHeader")
            doc.add_paragraph(f"Rol: {miembro.get('rol', 'No especificado')}")
            doc.add_paragraph(f"Experiencia: {miembro.get('experiencia', 'No especificada')}")

            skills = miembro.get('skills', [])
            if skills:
                doc.add_paragraph("Habilidades Clave:")
                for skill in skills:
                    doc.add_paragraph(f"• {skill}", style="List Bullet")
        
        # Guardar documento temporalmente
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_path = tmp_file.name
            doc.save(tmp_path)
            
        file_name = f"Propuesta_{proposal_data.get('cliente', 'Cliente')}.docx"
        
        return FileResponse(
            path=tmp_path,
            filename=file_name,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error al generar documento: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error al generar el documento: {str(e)}"
        )
