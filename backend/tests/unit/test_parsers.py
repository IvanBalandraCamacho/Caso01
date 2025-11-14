
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Protocol
import logging
import mimetypes

import pandas as pd
import PyPDF2
import docx
import pdfplumber  # Más robusto que PyPDF2
from core.exceptions import FileProcessingException

logger = logging.getLogger(__name__)

# ============================================================================
# PROTOCOLO PARA PARSERS
# ============================================================================

class FileParser(Protocol):
    """Protocolo que define la interfaz de un parser"""
    
    def can_parse(self, file_path: Path) -> bool:
        """Verifica si el parser puede procesar este archivo"""
        ...
    
    def extract_text(self, file_path: Path) -> str:
        """Extrae texto del archivo"""
        ...


# ============================================================================
# PARSERS ESPECÍFICOS
# ============================================================================

class PDFParser:
    """Parser para archivos PDF con fallback"""
    
    MIME_TYPES = ["application/pdf"]
    
    def can_parse(self, file_path: Path) -> bool:
        mime_type = mimetypes.guess_type(file_path)[0]
        return mime_type in self.MIME_TYPES
    
    def extract_text(self, file_path: Path) -> str:
        """Extrae texto usando pdfplumber (primero) o PyPDF2 (fallback)"""
        try:
            # Método 1: pdfplumber (mejor para PDFs complejos)
            text = self._extract_with_pdfplumber(file_path)
            if text and len(text.strip()) > 100:
                logger.info(f"PDF parseado con pdfplumber: {len(text)} caracteres")
                return text
        except Exception as e:
            logger.warning(f"pdfplumber falló, usando PyPDF2: {e}")
        
        # Método 2: PyPDF2 (fallback)
        try:
            text = self._extract_with_pypdf2(file_path)
            logger.info(f"PDF parseado con PyPDF2: {len(text)} caracteres")
            return text
        except Exception as e:
            raise FileProcessingException(file_path.name, f"Error en ambos parsers: {e}")
    
    def _extract_with_pdfplumber(self, file_path: Path) -> str:
        """Extracción con pdfplumber"""
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
    
    def _extract_with_pypdf2(self, file_path: Path) -> str:
        """Extracción con PyPDF2"""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text


class DOCXParser:
    """Parser para archivos Word"""
    
    MIME_TYPES = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]
    
    def can_parse(self, file_path: Path) -> bool:
        mime_type = mimetypes.guess_type(file_path)[0]
        return mime_type in self.MIME_TYPES or file_path.suffix == ".docx"
    
    def extract_text(self, file_path: Path) -> str:
        """Extrae texto de DOCX incluyendo tablas"""
        try:
            doc = docx.Document(file_path)
            
            # Extraer párrafos
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extraer tablas
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    tables_text.append(row_text)
            
            full_text = "\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n=== TABLAS ===\n" + "\n".join(tables_text)
            
            logger.info(f"DOCX parseado: {len(full_text)} caracteres")
            return full_text
            
        except Exception as e:
            raise FileProcessingException(file_path.name, str(e))


class ExcelParser:
    """Parser para archivos Excel"""
    
    MIME_TYPES = [
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel"
    ]
    
    def can_parse(self, file_path: Path) -> bool:
        mime_type = mimetypes.guess_type(file_path)[0]
        return (
            mime_type in self.MIME_TYPES 
            or file_path.suffix in [".xlsx", ".xls"]
        )
    
    def extract_text(self, file_path: Path) -> str:
        """Extrae texto de Excel (todas las hojas)"""
        try:
            # Leer todas las hojas
            excel_file = pd.ExcelFile(file_path)
            
            all_sheets_text = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Limpiar DataFrame
                df = df.dropna(how='all')  # Eliminar filas vacías
                
                sheet_text = f"\n=== Hoja: {sheet_name} ===\n"
                sheet_text += df.to_string(index=False)
                all_sheets_text.append(sheet_text)
            
            full_text = "\n\n".join(all_sheets_text)
            logger.info(f"Excel parseado: {len(full_text)} caracteres de {len(excel_file.sheet_names)} hojas")
            return full_text
            
        except Exception as e:
            raise FileProcessingException(file_path.name, str(e))


class CSVParser:
    """Parser para archivos CSV"""
    
    MIME_TYPES = ["text/csv", "application/csv"]
    
    def can_parse(self, file_path: Path) -> bool:
        mime_type = mimetypes.guess_type(file_path)[0]
        return mime_type in self.MIME_TYPES or file_path.suffix == ".csv"
    
    def extract_text(self, file_path: Path) -> str:
        """Extrae texto de CSV con detección automática de delimitadores"""
        try:
            # Pandas detecta automáticamente el delimitador
            df = pd.read_csv(file_path, encoding='utf-8')
            
            # Limpiar
            df = df.dropna(how='all')
            
            text = df.to_string(index=False)
            logger.info(f"CSV parseado: {len(text)} caracteres, {len(df)} filas")
            return text
            
        except UnicodeDecodeError:
            # Intentar con otra codificación
            df = pd.read_csv(file_path, encoding='latin-1')
            text = df.to_string(index=False)
            logger.warning(f"CSV parseado con encoding latin-1")
            return text
            
        except Exception as e:
            raise FileProcessingException(file_path.name, str(e))


class TextParser:
    """Parser para archivos de texto plano"""
    
    MIME_TYPES = ["text/plain"]
    
    def can_parse(self, file_path: Path) -> bool:
        mime_type = mimetypes.guess_type(file_path)[0]
        return mime_type in self.MIME_TYPES or file_path.suffix == ".txt"
    
    def extract_text(self, file_path: Path) -> str:
        """Lee archivo de texto con múltiples encodings"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                    logger.info(f"Texto plano leído con {encoding}: {len(text)} caracteres")
                    return text
            except UnicodeDecodeError:
                continue
        
        raise FileProcessingException(
            file_path.name, 
            f"No se pudo leer con ninguno de estos encodings: {encodings}"
        )


# ============================================================================
# FACTORY DE PARSERS
# ============================================================================

class ParserFactory:
    """Factory para seleccionar el parser correcto"""
    
    def __init__(self):
        self.parsers: list[FileParser] = [
            PDFParser(),
            DOCXParser(),
            ExcelParser(),
            CSVParser(),
            TextParser(),
        ]
    
    def get_parser(self, file_path: Path) -> FileParser:
        """Selecciona el parser apropiado para el archivo"""
        for parser in self.parsers:
            if parser.can_parse(file_path):
                logger.debug(f"Parser seleccionado: {parser.__class__.__name__}")
                return parser
        
        # Ningún parser puede procesarlo
        raise FileProcessingException(
            file_path.name,
            f"Tipo de archivo no soportado: {mimetypes.guess_type(file_path)[0]}"
        )


# ============================================================================
# API PÚBLICA
# ============================================================================

# Instancia global del factory
_parser_factory = ParserFactory()

def extract_text_from_file(file_path: Path) -> str:
    """
    Función principal para extraer texto de cualquier archivo soportado.
    
    Args:
        file_path: Ruta al archivo
    
    Returns:
        Texto extraído del archivo
    
    Raises:
        FileProcessingException: Si el archivo no se puede procesar
    """
    if not file_path.exists():
        raise FileProcessingException(file_path.name, "Archivo no existe")
    
    if file_path.stat().st_size == 0:
        raise FileProcessingException(file_path.name, "Archivo vacío")
    
    logger.info(f"Extrayendo texto de: {file_path.name}")
    
    try:
        parser = _parser_factory.get_parser(file_path)
        text = parser.extract_text(file_path)
        
        if not text or len(text.strip()) < 10:
            raise FileProcessingException(
                file_path.name, 
                "El archivo no contiene texto suficiente"
            )
        
        return text
        
    except FileProcessingException:
        raise
    except Exception as e:
        logger.exception(f"Error inesperado al procesar {file_path.name}")
        raise FileProcessingException(file_path.name, str(e))


# ============================================================================
# VALIDACIONES ADICIONALES
# ============================================================================

def validate_file_for_processing(file_path: Path, max_size_mb: int = 50) -> None:
    """
    Valida que un archivo pueda ser procesado.
    
    Args:
        file_path: Ruta al archivo
        max_size_mb: Tamaño máximo en MB
    
    Raises:
        FileProcessingException: Si la validación falla
    """
    if not file_path.exists():
        raise FileProcessingException(file_path.name, "Archivo no encontrado")
    
    # Validar tamaño
    size_mb = file_path.stat().st_size / (1024 * 1024)
    if size_mb > max_size_mb:
        raise FileProcessingException(
            file_path.name,
            f"Archivo muy grande: {size_mb:.2f}MB (máximo: {max_size_mb}MB)"
        )
    
    # Validar que no sea un archivo ejecutable o peligroso
    dangerous_extensions = ['.exe', '.dll', '.so', '.sh', '.bat', '.cmd']
    if file_path.suffix.lower() in dangerous_extensions:
        raise FileProcessingException(
            file_path.name,
            f"Tipo de archivo peligroso: {file_path.suffix}"
        )
    
    logger.debug(f"Archivo validado: {file_path.name} ({size_mb:.2f}MB)")